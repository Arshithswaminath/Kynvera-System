"""
Text extraction for Knowledge Base uploads (txt, md, docx, pdf).
Fails soft: returns whatever text it can, '' on error.
"""
import logging
import os

logger = logging.getLogger(__name__)

MAX_TEXT_CHARS = 50000
ALLOWED_EXTENSIONS = {'txt', 'md', 'docx', 'pdf'}


def _clean(text: str) -> str:
    if not text:
        return ''
    # Strip null/control chars except common whitespace
    cleaned = ''.join(ch for ch in text if ch == '\n' or ch == '\t' or ch >= ' ')
    cleaned = cleaned.strip()
    if len(cleaned) > MAX_TEXT_CHARS:
        cleaned = cleaned[:MAX_TEXT_CHARS]
    return cleaned


def _extract_txt(path: str) -> str:
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        logger.warning('KB extract txt failed for %s: %s', path, e)
        return ''


def _extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells if c.text]
                if cells:
                    parts.append(' | '.join(cells))
        return '\n'.join(parts)
    except Exception as e:
        logger.warning('KB extract docx failed for %s: %s', path, e)
        return ''


def _extract_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or '')
            except Exception:
                continue
        return '\n'.join(parts)
    except Exception as e:
        logger.warning('KB extract pdf failed for %s: %s', path, e)
        return ''


def extract_text(stored_path: str, file_type: str) -> str:
    """Extract plain text from a stored file by type."""
    if not stored_path or not os.path.isfile(stored_path):
        return ''
    ext = (file_type or '').lower().lstrip('.')
    if not ext:
        ext = (os.path.splitext(stored_path)[1] or '').lower().lstrip('.')

    if ext in ('txt', 'md'):
        raw = _extract_txt(stored_path)
    elif ext == 'docx':
        raw = _extract_docx(stored_path)
    elif ext == 'pdf':
        raw = _extract_pdf(stored_path)
    else:
        raw = ''
    return _clean(raw)
